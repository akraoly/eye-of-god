"""
ReportGeneratorService — génère des rapports d'audit professionnels
aux formats PDF, HTML, DOCX et Markdown.
"""
from __future__ import annotations

import asyncio
import json
import logging
import os
import uuid
from datetime import datetime
from pathlib import Path

log = logging.getLogger(__name__)

_REPORTS_DIR = "./data/reports"


# ── Template HTML ─────────────────────────────────────────────────────────────

def _html_template(data: dict, company_name: str = "AEGIS AI") -> str:
    """Construit le rapport HTML complet depuis les données collectées."""
    campaign_id = data.get("campaign_id", "N/A")
    generated_at = data.get("generated_at", datetime.utcnow().isoformat())
    summary_text = data.get("_summary", "Aucun résumé disponible.")
    risk_score = data.get("_risk_score", 0)
    title = data.get("_title", f"Rapport d'Audit — Campagne {campaign_id}")

    risk_color = (
        "#ef4444" if risk_score >= 75
        else "#f97316" if risk_score >= 50
        else "#eab308" if risk_score >= 25
        else "#22c55e"
    )
    risk_label = (
        "CRITIQUE" if risk_score >= 75
        else "ÉLEVÉ" if risk_score >= 50
        else "MODÉRÉ" if risk_score >= 25
        else "FAIBLE"
    )

    # ── Sections ──────────────────────────────────────────────────────────────

    def badge(severity: str) -> str:
        s = severity.upper()
        css = {"CRITICAL": "critical", "HIGH": "high", "MEDIUM": "medium",
               "LOW": "low", "ÉLEVÉ": "high", "CRITIQUE": "critical"}.get(s, "low")
        return f'<span class="badge-{css}">{severity}</span>'

    def rows_table(headers: list, rows: list) -> str:
        ths = "".join(f"<th>{h}</th>" for h in headers)
        trs = ""
        for row in rows:
            tds = "".join(f"<td>{v}</td>" for v in row)
            trs += f"<tr>{tds}</tr>"
        return f"<table><thead><tr>{ths}</tr></thead><tbody>{trs}</tbody></table>"

    # Section screenshots
    screenshots = data.get("screenshots", [])
    shots_html = ""
    if screenshots:
        shot_items = "".join(
            f'<div class="screenshot-item"><div class="screenshot-meta">'
            f'<strong>{Path(s.get("file_path", "?")).name}</strong> '
            f'— {s.get("source", "")} — {s.get("captured_at", "")[:19]}</div></div>'
            for s in screenshots[:20]
        )
        shots_html = f"""
        <div class="section">
          <div class="section-title">Captures d'écran ({len(screenshots)})</div>
          <div class="screenshot-gallery">{shot_items}</div>
        </div>"""

    # Section audio
    audio = data.get("audio_recordings", [])
    audio_html = ""
    if audio:
        audio_rows = [(
            r.get("recording_id", "")[:8],
            r.get("target_id", "N/A"),
            r.get("mic_name", "N/A"),
            f'{r.get("duration", 0)}s',
            r.get("format", "wav"),
            r.get("created_at", "")[:19],
        ) for r in audio[:50]]
        audio_html = f"""
        <div class="section">
          <div class="section-title">Enregistrements Audio ({len(audio)})</div>
          {rows_table(["ID", "Cible", "Microphone", "Durée", "Format", "Date"], audio_rows)}
        </div>"""

    # Section caméras
    cameras = data.get("camera_snapshots", [])
    cam_html = ""
    if cameras:
        cam_rows = [(
            s.get("snapshot_id", "")[:8],
            s.get("camera_id", "N/A"),
            Path(s.get("file_path", "?")).name if s.get("file_path") else "N/A",
            s.get("taken_at", "")[:19],
        ) for s in cameras[:50]]
        cam_html = f"""
        <div class="section">
          <div class="section-title">Snapshots Caméras IP ({len(cameras)})</div>
          {rows_table(["ID", "Caméra ID", "Fichier", "Date"], cam_rows)}
        </div>"""

    # Section keystrokes
    keystrokes = data.get("keystrokes", [])
    ks_html = ""
    if keystrokes:
        ks_rows = [(
            k.get("target_id", "N/A"),
            k.get("app_name", "N/A"),
            k.get("window_title", "N/A"),
            (k.get("keystrokes", "") or "")[:60] + "...",
            k.get("captured_at", "")[:19],
        ) for k in keystrokes[:50]]
        ks_html = f"""
        <div class="section">
          <div class="section-title">Frappes Clavier Capturées ({len(keystrokes)})</div>
          {rows_table(["Cible", "Application", "Fenêtre", "Frappes (extrait)", "Date"], ks_rows)}
        </div>"""

    # Section formulaires capturés
    forms = data.get("captured_forms", [])
    forms_html = ""
    if forms:
        form_rows = [(
            f.get("target_id", "N/A"),
            (f.get("url", "") or "")[:60],
            json.dumps(f.get("form_data", {}))[:80],
            f.get("captured_at", "")[:19],
        ) for f in forms[:50]]
        forms_html = f"""
        <div class="section">
          <div class="section-title">Formulaires Capturés / Identifiants ({len(forms)})</div>
          {rows_table(["Cible", "URL", "Données (extrait)", "Date"], form_rows)}
        </div>"""

    # Section réseau
    network = data.get("network_captures", [])
    net_html = ""
    if network:
        net_rows = [(
            c.get("interface", "N/A"),
            c.get("packet_count", 0),
            c.get("creds_found", 0),
            c.get("status", "N/A"),
            c.get("started_at", "")[:19],
        ) for c in network[:50]]
        net_html = f"""
        <div class="section">
          <div class="section-title">Captures Réseau ({len(network)})</div>
          {rows_table(["Interface", "Paquets", "Credentials trouvés", "Statut", "Début"], net_rows)}
        </div>"""

    # Section BLE
    ble = data.get("ble_devices", [])
    ble_html = ""
    if ble:
        ble_rows = [(
            d.get("mac_address", "N/A"),
            d.get("name", "N/A"),
            d.get("manufacturer", "N/A"),
            d.get("device_type", "N/A"),
            "OUI" if d.get("is_tracker") else "NON",
            f'{d.get("rssi", 0)} dBm',
        ) for d in ble[:50]]
        ble_html = f"""
        <div class="section">
          <div class="section-title">Appareils BLE Découverts ({len(ble)})</div>
          {rows_table(["MAC", "Nom", "Fabricant", "Type", "Tracker?", "RSSI"], ble_rows)}
        </div>"""

    # Section RFID
    rfid = data.get("rfid_cards", [])
    rfid_html = ""
    if rfid:
        rfid_rows = [tuple(str(v)[:40] for v in list(r.values())[:5]) for r in rfid[:30]]
        headers_rfid = list(rfid[0].keys())[:5] if rfid else []
        rfid_html = f"""
        <div class="section">
          <div class="section-title">Cartes RFID Capturées ({len(rfid)})</div>
          {rows_table(headers_rfid, rfid_rows)}
        </div>"""

    # Section SDR
    sdr = data.get("sdr_recordings", [])
    sdr_html = ""
    if sdr:
        sdr_rows = [(
            f'{r.get("frequency_mhz", 0)} MHz',
            r.get("modulation", "N/A"),
            r.get("protocol", "N/A"),
            f'{r.get("duration", 0)}s',
            r.get("replay_count", 0),
            r.get("created_at", "")[:19],
        ) for r in sdr[:50]]
        sdr_html = f"""
        <div class="section">
          <div class="section-title">Enregistrements SDR ({len(sdr)})</div>
          {rows_table(["Fréquence", "Modulation", "Protocole", "Durée", "Replays", "Date"], sdr_rows)}
        </div>"""

    # Section MITRE ATT&CK
    mitre = data.get("mitre_stats", {})
    mitre_html = ""
    if mitre:
        mitre_events = mitre.get("events", [])
        evt_rows = [(
            e.get("technique_id", "N/A"),
            e.get("tactic_id", "N/A"),
            e.get("action_type", "N/A"),
            e.get("score", 0),
            "OUI" if e.get("success") else "NON",
            e.get("timestamp", "")[:19],
        ) for e in mitre_events[:50]]
        mitre_table = rows_table(
            ["Technique", "Tactique", "Action", "Score", "Succès", "Horodatage"],
            evt_rows
        ) if evt_rows else "<p>Aucun événement MITRE enregistré.</p>"

        mitre_html = f"""
        <div class="section">
          <div class="section-title">MITRE ATT&amp;CK — Statistiques</div>
          <div class="mitre-stats">
            <div class="stat-box"><div class="stat-num">{mitre.get("total_techniques", 0)}</div><div>Techniques</div></div>
            <div class="stat-box"><div class="stat-num">{mitre.get("total_tactics", 0)}</div><div>Tactiques</div></div>
            <div class="stat-box"><div class="stat-num">{mitre.get("total_score", 0)}</div><div>Score Total</div></div>
            <div class="stat-box"><div class="stat-num">{mitre.get("coverage", 0):.1f}%</div><div>Couverture</div></div>
          </div>
          {mitre_table}
        </div>"""

    # Table des matières
    toc_items = []
    toc_map = [
        ("1", "Résumé Exécutif"),
        ("2", "Score de Risque"),
    ]
    n = 3
    if screenshots:
        toc_map.append((str(n), f"Captures d'écran ({len(screenshots)})"))
        n += 1
    if audio:
        toc_map.append((str(n), f"Enregistrements Audio ({len(audio)})"))
        n += 1
    if cameras:
        toc_map.append((str(n), f"Snapshots Caméras ({len(cameras)})"))
        n += 1
    if keystrokes:
        toc_map.append((str(n), f"Frappes Clavier ({len(keystrokes)})"))
        n += 1
    if forms:
        toc_map.append((str(n), f"Formulaires / Identifiants ({len(forms)})"))
        n += 1
    if network:
        toc_map.append((str(n), f"Captures Réseau ({len(network)})"))
        n += 1
    if ble:
        toc_map.append((str(n), f"Appareils BLE ({len(ble)})"))
        n += 1
    if rfid:
        toc_map.append((str(n), f"Cartes RFID ({len(rfid)})"))
        n += 1
    if sdr:
        toc_map.append((str(n), f"Enregistrements SDR ({len(sdr)})"))
        n += 1
    if mitre:
        toc_map.append((str(n), "MITRE ATT&CK"))

    toc_html = "".join(
        f'<div class="toc-item"><span class="toc-num">{num}.</span> {label}</div>'
        for num, label in toc_map
    )

    return f"""<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0"/>
<title>{title}</title>
<style>
  body {{ font-family: Arial, sans-serif; color: #1a1a2e; background: white; margin: 0; padding: 0; }}
  .cover {{ text-align: center; padding: 100px 40px; border-bottom: 3px solid #1a237e; background: linear-gradient(135deg, #f8fafc 0%, #e8eaf6 100%); }}
  .cover-company {{ font-size: 14px; letter-spacing: 4px; color: #1a237e; text-transform: uppercase; margin-bottom: 40px; }}
  .title {{ font-size: 28px; font-weight: bold; color: #1a237e; margin: 20px 0; }}
  .cover-sub {{ font-size: 14px; color: #555; margin-top: 10px; }}
  .cover-meta {{ margin-top: 40px; font-size: 12px; color: #888; }}
  .cover-badge {{ display: inline-block; background: #1a237e; color: white; padding: 6px 20px; border-radius: 20px; font-size: 12px; letter-spacing: 2px; margin-top: 20px; }}
  .content {{ max-width: 1100px; margin: 0 auto; padding: 40px 40px; }}
  .section {{ margin: 40px 0; page-break-inside: avoid; }}
  .section-title {{ font-size: 20px; color: #1a237e; border-left: 4px solid #1a237e; padding-left: 12px; margin-bottom: 16px; font-weight: bold; }}
  .toc {{ background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 8px; padding: 24px; margin: 40px 0; }}
  .toc-title {{ font-size: 16px; font-weight: bold; color: #1a237e; margin-bottom: 16px; }}
  .toc-item {{ padding: 6px 0; border-bottom: 1px dotted #cbd5e0; font-size: 14px; }}
  .toc-num {{ font-weight: bold; color: #1a237e; margin-right: 8px; }}
  .badge-critical {{ background: #ef4444; color: white; padding: 2px 8px; border-radius: 4px; font-size: 12px; font-weight: bold; }}
  .badge-high {{ background: #f97316; color: white; padding: 2px 8px; border-radius: 4px; font-size: 12px; font-weight: bold; }}
  .badge-medium {{ background: #eab308; color: white; padding: 2px 8px; border-radius: 4px; font-size: 12px; font-weight: bold; }}
  .badge-low {{ background: #22c55e; color: white; padding: 2px 8px; border-radius: 4px; font-size: 12px; font-weight: bold; }}
  table {{ width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 13px; }}
  th {{ background: #1a237e; color: white; padding: 8px 12px; text-align: left; }}
  td {{ padding: 8px 12px; border-bottom: 1px solid #e2e8f0; word-break: break-all; }}
  tr:nth-child(even) {{ background: #f8fafc; }}
  .risk-block {{ text-align: center; padding: 40px; background: #f8fafc; border-radius: 12px; border: 2px solid {risk_color}; margin: 20px 0; }}
  .risk-score {{ font-size: 72px; font-weight: bold; color: {risk_color}; line-height: 1; }}
  .risk-label {{ font-size: 18px; font-weight: bold; color: {risk_color}; margin-top: 8px; }}
  .risk-bar-bg {{ background: #e2e8f0; border-radius: 4px; height: 12px; margin: 16px auto; max-width: 300px; }}
  .risk-bar-fill {{ background: {risk_color}; border-radius: 4px; height: 12px; width: {min(risk_score, 100):.0f}%; }}
  .summary-box {{ background: #f0f4ff; border-left: 4px solid #1a237e; padding: 20px 24px; border-radius: 0 8px 8px 0; font-size: 14px; line-height: 1.7; margin: 20px 0; white-space: pre-line; }}
  .screenshot-gallery {{ display: flex; flex-wrap: wrap; gap: 12px; }}
  .screenshot-item {{ background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 6px; padding: 10px; font-size: 12px; flex: 0 0 calc(50% - 6px); }}
  .screenshot-meta {{ color: #555; word-break: break-all; }}
  .mitre-stats {{ display: flex; gap: 20px; margin: 20px 0; flex-wrap: wrap; }}
  .stat-box {{ background: #f0f4ff; border: 1px solid #c7d2fe; border-radius: 8px; padding: 16px 24px; text-align: center; flex: 1; min-width: 100px; }}
  .stat-num {{ font-size: 32px; font-weight: bold; color: #1a237e; }}
  .footer {{ text-align: center; color: #666; font-size: 11px; border-top: 1px solid #e2e8f0; padding: 24px; margin-top: 60px; background: #f8fafc; }}
  @media print {{
    .cover {{ page-break-after: always; }}
    .section {{ page-break-inside: avoid; }}
  }}
</style>
</head>
<body>

<!-- PAGE TITRE -->
<div class="cover">
  <div class="cover-company">{company_name}</div>
  <div class="title">{title}</div>
  <div class="cover-sub">Rapport d'Audit de Sécurité Professionnel</div>
  <div class="cover-meta">
    Campagne ID: <strong>{campaign_id}</strong><br/>
    Généré le: <strong>{generated_at[:19].replace("T", " ")} UTC</strong>
  </div>
  <div class="cover-badge">CONFIDENTIEL</div>
</div>

<div class="content">

<!-- TABLE DES MATIÈRES -->
<div class="toc">
  <div class="toc-title">Table des Matières</div>
  {toc_html}
</div>

<!-- RÉSUMÉ EXÉCUTIF -->
<div class="section">
  <div class="section-title">1. Résumé Exécutif</div>
  <div class="summary-box">{summary_text}</div>
</div>

<!-- SCORE DE RISQUE -->
<div class="section">
  <div class="section-title">2. Score de Risque Global</div>
  <div class="risk-block">
    <div class="risk-score">{risk_score:.0f}</div>
    <div class="risk-label">{risk_label}</div>
    <div class="risk-bar-bg"><div class="risk-bar-fill"></div></div>
    <div style="font-size:12px;color:#888;margin-top:8px;">Score sur 100 — basé sur les vulnérabilités critiques, credentials capturés, accès C2</div>
  </div>
</div>

{shots_html}
{audio_html}
{cam_html}
{ks_html}
{forms_html}
{net_html}
{ble_html}
{rfid_html}
{sdr_html}
{mitre_html}

</div><!-- /content -->

<div class="footer">
  Rapport généré par <strong>{company_name} — L'Œil de Dieu</strong> |
  Campagne: {campaign_id} |
  {generated_at[:19].replace("T", " ")} UTC<br/>
  <em>Ce document est confidentiel. Distribution restreinte aux personnes autorisées.</em>
</div>
</body>
</html>"""


# ── Service principal ─────────────────────────────────────────────────────────

class ReportGeneratorService:
    """Génère des rapports d'audit aux formats PDF, HTML, DOCX et Markdown."""

    def __init__(self):
        os.makedirs(_REPORTS_DIR, exist_ok=True)

    # ── API publique ──────────────────────────────────────────────────────────

    async def generate_report(
        self,
        campaign_id: str,
        format: str = "pdf",
        options: dict | None = None,
        title: str | None = None,
        db=None,
        generated_by: str | None = None,
    ) -> dict:
        """Collecte les données, génère le rapport, sauvegarde en DB."""
        options = options or {}
        report_id = str(uuid.uuid4())
        title = title or f"Rapport d'Audit — {campaign_id}"
        fmt = format.lower()

        # 1. Créer l'entrée DB en statut "generating"
        if db:
            try:
                from database.models_reporting import AuditReport
                entry = AuditReport(
                    report_id=report_id,
                    campaign_id=campaign_id,
                    title=title,
                    format=fmt,
                    status="generating",
                    options=options,
                    generated_by=generated_by,
                )
                db.add(entry)
                db.commit()
            except Exception as e:
                log.warning("DB insert error: %s", e)

        # 2. Collecter les données
        try:
            data = await self._collect_all_data(campaign_id, db)
        except Exception as e:
            log.error("Data collection failed: %s", e)
            data = {
                "campaign_id": campaign_id,
                "generated_at": datetime.utcnow().isoformat(),
                "screenshots": [],
                "audio_recordings": [],
                "camera_snapshots": [],
                "keystrokes": [],
                "captured_forms": [],
                "network_captures": [],
                "ble_devices": [],
                "rfid_cards": [],
                "sdr_recordings": [],
                "mitre_stats": {},
            }

        # 3. Enrichir les métadonnées
        data["_title"] = title
        data["_summary"] = await self._build_summary(data)
        data["_risk_score"] = await self._calculate_risk_score(data)

        # 4. Répertoire de sortie
        output_dir = Path(_REPORTS_DIR) / campaign_id
        output_dir.mkdir(parents=True, exist_ok=True)
        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        base_name = f"audit_{campaign_id}_{ts}"

        # 5. Générer selon le format
        file_path: str | None = None
        actual_fmt = fmt
        try:
            if fmt == "pdf":
                file_path = await self.generate_pdf(data, str(output_dir / f"{base_name}.pdf"))
                if file_path and file_path.endswith(".html"):
                    actual_fmt = "html"
            elif fmt == "html":
                file_path = await self.generate_html(data, str(output_dir / f"{base_name}.html"))
            elif fmt == "docx":
                file_path = await self.generate_docx(data, str(output_dir / f"{base_name}.docx"))
                if file_path and file_path.endswith(".json"):
                    actual_fmt = "json"
            elif fmt == "markdown":
                file_path = await self.generate_markdown(data, str(output_dir / f"{base_name}.md"))
            else:
                file_path = await self.generate_html(data, str(output_dir / f"{base_name}.html"))
                actual_fmt = "html"
        except Exception as e:
            log.error("Generation failed for format %s: %s", fmt, e)

        # 6. Taille du fichier
        file_size = None
        if file_path and Path(file_path).exists():
            file_size = Path(file_path).stat().st_size

        # 7. Mettre à jour l'entrée DB
        if db:
            try:
                from database.models_reporting import AuditReport
                entry = db.query(AuditReport).filter(
                    AuditReport.report_id == report_id
                ).first()
                if entry:
                    entry.file_path = file_path
                    entry.file_size = file_size
                    entry.format = actual_fmt
                    entry.status = "ready" if file_path else "error"
                    entry.summary = data.get("_summary", "")
                    entry.risk_score = data.get("_risk_score", 0)
                    db.commit()
            except Exception as e:
                log.warning("DB update error: %s", e)

        return {
            "report_id": report_id,
            "campaign_id": campaign_id,
            "title": title,
            "file_path": file_path,
            "format": actual_fmt,
            "file_size": file_size,
            "status": "ready" if file_path else "error",
            "risk_score": data.get("_risk_score", 0),
            "summary": data.get("_summary", ""),
        }

    # ── Collecte des données ──────────────────────────────────────────────────

    async def _collect_all_data(self, campaign_id: str, db) -> dict:
        from services.reporting.collectors.evidence_collector import EvidenceCollector
        collector = EvidenceCollector()
        return await collector.collect_all(campaign_id, db)

    # ── Résumé exécutif ───────────────────────────────────────────────────────

    async def _build_summary(self, data: dict) -> str:
        campaign_id = data.get("campaign_id", "N/A")
        generated_at = data.get("generated_at", "")[:19]
        screenshots = len(data.get("screenshots", []))
        audio = len(data.get("audio_recordings", []))
        cameras = len(data.get("camera_snapshots", []))
        keystrokes = len(data.get("keystrokes", []))
        forms = len(data.get("captured_forms", []))
        network = len(data.get("network_captures", []))
        ble = len(data.get("ble_devices", []))
        rfid = len(data.get("rfid_cards", []))
        sdr = len(data.get("sdr_recordings", []))
        mitre = data.get("mitre_stats", {})
        techniques = mitre.get("total_techniques", 0)
        tactics = mitre.get("total_tactics", 0)
        risk = data.get("_risk_score", 0)

        level = (
            "CRITIQUE — intervention immédiate requise"
            if risk >= 75 else
            "ÉLEVÉ — mesures correctives urgentes"
            if risk >= 50 else
            "MODÉRÉ — améliorations recommandées"
            if risk >= 25 else
            "FAIBLE — posture de sécurité acceptable"
        )

        lines = [
            f"Campagne d'audit #{campaign_id} — Généré le {generated_at} UTC",
            f"Niveau de risque global : {risk:.0f}/100 ({level})",
            "",
            "Preuves collectées :",
        ]
        if screenshots:
            lines.append(f"  • {screenshots} capture(s) d'écran")
        if audio:
            lines.append(f"  • {audio} enregistrement(s) audio")
        if cameras:
            lines.append(f"  • {cameras} snapshot(s) caméra IP")
        if keystrokes:
            lines.append(f"  • {keystrokes} session(s) de keylogging")
        if forms:
            lines.append(f"  • {forms} formulaire(s)/identifiant(s) capturé(s)")
        if network:
            lines.append(f"  • {network} capture(s) réseau")
        if ble:
            lines.append(f"  • {ble} appareil(s) BLE découvert(s)")
        if rfid:
            lines.append(f"  • {rfid} carte(s) RFID capturée(s)")
        if sdr:
            lines.append(f"  • {sdr} enregistrement(s) SDR")
        if techniques or tactics:
            lines.append(f"  • {techniques} technique(s) MITRE ATT&CK — {tactics} tactique(s)")

        total = screenshots + audio + cameras + keystrokes + forms + network + ble + rfid + sdr
        if total == 0:
            lines.append("  • Aucune preuve collectée pour cette campagne.")

        lines.extend([
            "",
            "Ce rapport présente l'ensemble des preuves collectées lors de l'opération d'audit.",
            "Les données sont classifiées CONFIDENTIEL et destinées aux équipes autorisées uniquement.",
        ])
        return "\n".join(lines)

    # ── Score de risque ───────────────────────────────────────────────────────

    async def _calculate_risk_score(self, data: dict) -> float:
        score = 0.0

        # Credentials capturés (haute criticité)
        forms = data.get("captured_forms", [])
        if forms:
            score += min(len(forms) * 8, 30)

        keystrokes = data.get("keystrokes", [])
        if keystrokes:
            score += min(len(keystrokes) * 3, 15)

        # Accès C2 via MITRE
        mitre = data.get("mitre_stats", {})
        techniques = mitre.get("total_techniques", 0)
        tactics = mitre.get("total_tactics", 0)
        if techniques:
            score += min(techniques * 1.5, 20)
        if tactics:
            score += min(tactics * 2, 15)

        # Preuves d'accès physique
        rfid = data.get("rfid_cards", [])
        if rfid:
            score += min(len(rfid) * 5, 15)

        ble = data.get("ble_devices", [])
        trackers = [d for d in ble if d.get("is_tracker")]
        if trackers:
            score += min(len(trackers) * 3, 10)

        # Captures réseau avec credentials
        network = data.get("network_captures", [])
        creds_found = sum(c.get("creds_found", 0) for c in network)
        if creds_found:
            score += min(creds_found * 4, 15)

        # Audio/vidéo — surveillance
        audio = data.get("audio_recordings", [])
        cameras = data.get("camera_snapshots", [])
        if audio:
            score += min(len(audio) * 2, 10)
        if cameras:
            score += min(len(cameras) * 1, 5)

        return min(round(score, 1), 100.0)

    # ── Générateurs de format ─────────────────────────────────────────────────

    async def generate_html(self, data: dict, output_path: str) -> str:
        """Génère un rapport HTML complet."""
        company_name = data.get("_company_name", "AEGIS AI")
        html_content = _html_template(data, company_name=company_name)
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        log.info("HTML report generated: %s", output_path)
        return output_path

    async def generate_pdf(self, data: dict, output_path: str) -> str:
        """Tente WeasyPrint, fallback wkhtmltopdf, fallback HTML."""
        html_path = output_path.replace(".pdf", ".html")
        await self.generate_html(data, html_path)

        # Tentative 1 — WeasyPrint
        try:
            import importlib
            weasyprint = importlib.import_module("weasyprint")
            weasyprint.HTML(filename=html_path).write_pdf(output_path)
            log.info("PDF generated via WeasyPrint: %s", output_path)
            return output_path
        except ImportError:
            log.warning("WeasyPrint non disponible, tentative wkhtmltopdf…")
        except Exception as e:
            log.warning("WeasyPrint failed: %s, tentative wkhtmltopdf…", e)

        # Tentative 2 — wkhtmltopdf
        try:
            proc = await asyncio.create_subprocess_exec(
                "wkhtmltopdf",
                "--quiet",
                "--enable-local-file-access",
                html_path,
                output_path,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            _, stderr = await asyncio.wait_for(proc.communicate(), timeout=60)
            if proc.returncode == 0 and Path(output_path).exists():
                log.info("PDF generated via wkhtmltopdf: %s", output_path)
                return output_path
            else:
                log.warning("wkhtmltopdf failed (rc=%d): %s", proc.returncode,
                            stderr.decode(errors="replace"))
        except FileNotFoundError:
            log.warning("wkhtmltopdf non installé")
        except asyncio.TimeoutError:
            log.warning("wkhtmltopdf timeout")
        except Exception as e:
            log.warning("wkhtmltopdf error: %s", e)

        # Fallback — retourner le HTML
        log.warning("PDF indisponible — le rapport sera livré en HTML: %s", html_path)
        return html_path

    async def generate_docx(self, data: dict, output_path: str) -> str:
        """Tente python-docx, fallback JSON."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Page de titre
            title_para = doc.add_heading(data.get("_title", "Rapport d'Audit"), 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f"Campagne: {data.get('campaign_id', 'N/A')}")
            doc.add_paragraph(f"Généré: {data.get('generated_at', '')[:19]}")
            doc.add_paragraph(f"Score de risque: {data.get('_risk_score', 0):.0f}/100")
            doc.add_page_break()

            # Résumé exécutif
            doc.add_heading("Résumé Exécutif", 1)
            doc.add_paragraph(data.get("_summary", ""))
            doc.add_page_break()

            # Sections
            sections_map = {
                "Captures d'écran": data.get("screenshots", []),
                "Enregistrements Audio": data.get("audio_recordings", []),
                "Snapshots Caméras": data.get("camera_snapshots", []),
                "Frappes Clavier": data.get("keystrokes", []),
                "Formulaires Capturés": data.get("captured_forms", []),
                "Captures Réseau": data.get("network_captures", []),
                "Appareils BLE": data.get("ble_devices", []),
                "Cartes RFID": data.get("rfid_cards", []),
                "Enregistrements SDR": data.get("sdr_recordings", []),
            }
            for section_title, items in sections_map.items():
                if items:
                    doc.add_heading(section_title, 2)
                    doc.add_paragraph(f"{len(items)} élément(s) collecté(s).")
                    if items and isinstance(items[0], dict):
                        keys = list(items[0].keys())[:6]
                        table = doc.add_table(rows=1, cols=len(keys))
                        table.style = "Table Grid"
                        hdr = table.rows[0].cells
                        for i, k in enumerate(keys):
                            hdr[i].text = str(k)
                        for item in items[:30]:
                            row = table.add_row().cells
                            for i, k in enumerate(keys):
                                row[i].text = str(item.get(k, ""))[:100]

            # MITRE
            mitre = data.get("mitre_stats", {})
            if mitre:
                doc.add_heading("MITRE ATT&CK", 2)
                doc.add_paragraph(
                    f"Techniques: {mitre.get('total_techniques', 0)} | "
                    f"Tactiques: {mitre.get('total_tactics', 0)} | "
                    f"Score: {mitre.get('total_score', 0)}"
                )

            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            log.info("DOCX generated: %s", output_path)
            return output_path

        except ImportError:
            log.warning("python-docx non disponible — fallback JSON")
        except Exception as e:
            log.warning("DOCX generation failed: %s — fallback JSON", e)

        # Fallback JSON
        json_path = output_path.replace(".docx", ".json")
        Path(json_path).parent.mkdir(parents=True, exist_ok=True)
        payload = {
            "report_id": data.get("_title", ""),
            "campaign_id": data.get("campaign_id", ""),
            "generated_at": data.get("generated_at", ""),
            "risk_score": data.get("_risk_score", 0),
            "summary": data.get("_summary", ""),
            "counts": {
                "screenshots": len(data.get("screenshots", [])),
                "audio_recordings": len(data.get("audio_recordings", [])),
                "camera_snapshots": len(data.get("camera_snapshots", [])),
                "keystrokes": len(data.get("keystrokes", [])),
                "captured_forms": len(data.get("captured_forms", [])),
                "network_captures": len(data.get("network_captures", [])),
                "ble_devices": len(data.get("ble_devices", [])),
                "rfid_cards": len(data.get("rfid_cards", [])),
                "sdr_recordings": len(data.get("sdr_recordings", [])),
            },
        }
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        return json_path

    async def generate_markdown(self, data: dict, output_path: str) -> str:
        """Génère un rapport Markdown complet (pas de dépendance externe)."""
        campaign_id = data.get("campaign_id", "N/A")
        generated_at = data.get("generated_at", "")[:19]
        title = data.get("_title", f"Rapport d'Audit — {campaign_id}")
        risk_score = data.get("_risk_score", 0)
        summary = data.get("_summary", "")

        lines = [
            f"# {title}",
            "",
            f"> **Campagne:** `{campaign_id}`  ",
            f"> **Généré:** {generated_at} UTC  ",
            f"> **Classification:** CONFIDENTIEL",
            "",
            "---",
            "",
            "## Table des Matières",
            "",
            "1. [Résumé Exécutif](#résumé-exécutif)",
            "2. [Score de Risque](#score-de-risque)",
            "3. [Captures d'écran](#captures-décran)",
            "4. [Enregistrements Audio](#enregistrements-audio)",
            "5. [Snapshots Caméras](#snapshots-caméras)",
            "6. [Frappes Clavier](#frappes-clavier)",
            "7. [Formulaires Capturés](#formulaires-capturés)",
            "8. [Captures Réseau](#captures-réseau)",
            "9. [Appareils BLE](#appareils-ble)",
            "10. [Cartes RFID](#cartes-rfid)",
            "11. [Enregistrements SDR](#enregistrements-sdr)",
            "12. [MITRE ATT&CK](#mitre-attck)",
            "",
            "---",
            "",
            "## Résumé Exécutif",
            "",
            summary,
            "",
            "---",
            "",
            "## Score de Risque",
            "",
            f"**Score global: {risk_score:.0f} / 100**",
            "",
            f"```",
            f"{'█' * int(risk_score // 5)}{'░' * (20 - int(risk_score // 5))} {risk_score:.0f}%",
            f"```",
            "",
            "---",
            "",
        ]

        def md_table(headers: list, rows: list) -> list[str]:
            result = []
            result.append("| " + " | ".join(str(h) for h in headers) + " |")
            result.append("|" + "|".join(["---"] * len(headers)) + "|")
            for row in rows[:30]:
                result.append("| " + " | ".join(str(v)[:60].replace("|", "\\|") for v in row) + " |")
            result.append("")
            return result

        # Screenshots
        screenshots = data.get("screenshots", [])
        lines.append("## Captures d'écran")
        lines.append("")
        if screenshots:
            lines.extend(md_table(
                ["Source", "Cible", "Fichier", "Date"],
                [(s.get("source",""), s.get("session_id",""), Path(s.get("file_path","?")).name, s.get("captured_at","")[:19])
                 for s in screenshots]
            ))
        else:
            lines.append("_Aucune capture d'écran._")
            lines.append("")
        lines.append("---")
        lines.append("")

        # Audio
        audio = data.get("audio_recordings", [])
        lines.append("## Enregistrements Audio")
        lines.append("")
        if audio:
            lines.extend(md_table(
                ["ID", "Cible", "Durée", "Format", "Date"],
                [(r.get("recording_id","")[:8], r.get("target_id","N/A"), f'{r.get("duration",0)}s', r.get("format","wav"), r.get("created_at","")[:19])
                 for r in audio]
            ))
        else:
            lines.append("_Aucun enregistrement audio._")
            lines.append("")
        lines.append("---")
        lines.append("")

        # Caméras
        cameras = data.get("camera_snapshots", [])
        lines.append("## Snapshots Caméras")
        lines.append("")
        if cameras:
            lines.extend(md_table(
                ["Snapshot ID", "Caméra ID", "Fichier", "Date"],
                [(s.get("snapshot_id","")[:8], s.get("camera_id","N/A"), Path(s.get("file_path","?")).name if s.get("file_path") else "N/A", s.get("taken_at","")[:19])
                 for s in cameras]
            ))
        else:
            lines.append("_Aucun snapshot caméra._")
            lines.append("")
        lines.append("---")
        lines.append("")

        # Keystrokes
        keystrokes = data.get("keystrokes", [])
        lines.append("## Frappes Clavier")
        lines.append("")
        if keystrokes:
            lines.extend(md_table(
                ["Cible", "App", "Fenêtre", "Extrait", "Date"],
                [(k.get("target_id","N/A"), k.get("app_name","N/A"), k.get("window_title","N/A"),
                  (k.get("keystrokes","") or "")[:40]+"…", k.get("captured_at","")[:19])
                 for k in keystrokes]
            ))
        else:
            lines.append("_Aucune frappe clavier._")
            lines.append("")
        lines.append("---")
        lines.append("")

        # Formulaires
        forms = data.get("captured_forms", [])
        lines.append("## Formulaires Capturés")
        lines.append("")
        if forms:
            lines.extend(md_table(
                ["Cible", "URL", "Données (extrait)", "Date"],
                [(f.get("target_id","N/A"), (f.get("url","") or "")[:50], json.dumps(f.get("form_data",{}))[:50], f.get("captured_at","")[:19])
                 for f in forms]
            ))
        else:
            lines.append("_Aucun formulaire capturé._")
            lines.append("")
        lines.append("---")
        lines.append("")

        # Réseau
        network = data.get("network_captures", [])
        lines.append("## Captures Réseau")
        lines.append("")
        if network:
            lines.extend(md_table(
                ["Interface", "Paquets", "Creds trouvés", "Statut", "Début"],
                [(c.get("interface","N/A"), c.get("packet_count",0), c.get("creds_found",0), c.get("status","N/A"), c.get("started_at","")[:19])
                 for c in network]
            ))
        else:
            lines.append("_Aucune capture réseau._")
            lines.append("")
        lines.append("---")
        lines.append("")

        # BLE
        ble = data.get("ble_devices", [])
        lines.append("## Appareils BLE")
        lines.append("")
        if ble:
            lines.extend(md_table(
                ["MAC", "Nom", "Fabricant", "Type", "Tracker?"],
                [(d.get("mac_address","N/A"), d.get("name","N/A"), d.get("manufacturer","N/A"), d.get("device_type","N/A"), "OUI" if d.get("is_tracker") else "NON")
                 for d in ble]
            ))
        else:
            lines.append("_Aucun appareil BLE._")
            lines.append("")
        lines.append("---")
        lines.append("")

        # RFID
        rfid = data.get("rfid_cards", [])
        lines.append("## Cartes RFID")
        lines.append("")
        if rfid:
            keys = list(rfid[0].keys())[:5]
            lines.extend(md_table(
                keys,
                [tuple(str(r.get(k,""))[:40] for k in keys) for r in rfid]
            ))
        else:
            lines.append("_Aucune carte RFID._")
            lines.append("")
        lines.append("---")
        lines.append("")

        # SDR
        sdr = data.get("sdr_recordings", [])
        lines.append("## Enregistrements SDR")
        lines.append("")
        if sdr:
            lines.extend(md_table(
                ["Fréquence (MHz)", "Modulation", "Protocole", "Durée", "Date"],
                [(r.get("frequency_mhz","N/A"), r.get("modulation","N/A"), r.get("protocol","N/A"), f'{r.get("duration",0)}s', r.get("created_at","")[:19])
                 for r in sdr]
            ))
        else:
            lines.append("_Aucun enregistrement SDR._")
            lines.append("")
        lines.append("---")
        lines.append("")

        # MITRE
        mitre = data.get("mitre_stats", {})
        lines.append("## MITRE ATT&CK")
        lines.append("")
        lines.append(f"| Techniques | Tactiques | Score | Couverture |")
        lines.append(f"|---|---|---|---|")
        lines.append(f"| {mitre.get('total_techniques',0)} | {mitre.get('total_tactics',0)} | {mitre.get('total_score',0)} | {mitre.get('coverage',0):.1f}% |")
        lines.append("")
        mitre_events = mitre.get("events", [])
        if mitre_events:
            lines.extend(md_table(
                ["Technique", "Tactique", "Action", "Score", "Succès", "Date"],
                [(e.get("technique_id",""), e.get("tactic_id",""), e.get("action_type",""), e.get("score",0), "OUI" if e.get("success") else "NON", e.get("timestamp","")[:19])
                 for e in mitre_events]
            ))
        lines.append("---")
        lines.append("")
        lines.append(f"*Rapport généré par AEGIS AI — L'Œil de Dieu | {generated_at} UTC*")

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        log.info("Markdown report generated: %s", output_path)
        return output_path

    # ── CRUD ──────────────────────────────────────────────────────────────────

    async def get_report_status(self, report_id: str, db) -> dict | None:
        try:
            from database.models_reporting import AuditReport
            entry = db.query(AuditReport).filter(
                AuditReport.report_id == report_id
            ).first()
            if not entry:
                return None
            return {
                "report_id": entry.report_id,
                "campaign_id": entry.campaign_id,
                "title": entry.title,
                "format": entry.format,
                "file_path": entry.file_path,
                "file_size": entry.file_size,
                "pages_count": entry.pages_count,
                "status": entry.status,
                "risk_score": entry.risk_score,
                "summary": entry.summary,
                "generated_by": entry.generated_by,
                "created_at": entry.created_at.isoformat() if entry.created_at else None,
            }
        except Exception as e:
            log.error("get_report_status error: %s", e)
            return None

    async def list_reports(self, campaign_id: str, db) -> list[dict]:
        try:
            from database.models_reporting import AuditReport
            entries = (
                db.query(AuditReport)
                .filter(AuditReport.campaign_id == campaign_id)
                .order_by(AuditReport.created_at.desc())
                .limit(100)
                .all()
            )
            return [
                {
                    "report_id": e.report_id,
                    "campaign_id": e.campaign_id,
                    "title": e.title,
                    "format": e.format,
                    "file_path": e.file_path,
                    "file_size": e.file_size,
                    "pages_count": e.pages_count,
                    "status": e.status,
                    "risk_score": e.risk_score,
                    "summary": e.summary,
                    "generated_by": e.generated_by,
                    "created_at": e.created_at.isoformat() if e.created_at else None,
                }
                for e in entries
            ]
        except Exception as e:
            log.error("list_reports error: %s", e)
            return []

    async def delete_report(self, report_id: str, db) -> bool:
        try:
            from database.models_reporting import AuditReport
            entry = db.query(AuditReport).filter(
                AuditReport.report_id == report_id
            ).first()
            if not entry:
                return False
            # Supprimer le fichier
            if entry.file_path and Path(entry.file_path).exists():
                try:
                    Path(entry.file_path).unlink()
                    # Supprimer aussi le .html si le report est PDF
                    html_path = entry.file_path.replace(".pdf", ".html")
                    if Path(html_path).exists():
                        Path(html_path).unlink()
                except Exception as fe:
                    log.warning("File deletion error: %s", fe)
            db.delete(entry)
            db.commit()
            return True
        except Exception as e:
            log.error("delete_report error: %s", e)
            return False
